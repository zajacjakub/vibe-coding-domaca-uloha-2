"""Convert audit report markdown to docx.

Handles only the subset of Markdown that the audit reports use:
- ATX headings (# .. ####)
- Bullet lists (- item)
- Pipe tables
- Blockquotes (> ...)
- Horizontal rules (---)
- Inline: **bold**, *italic*, `code`, [text](url)
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


INLINE_PATTERN = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\)"
    r"|(?<!\*)\*(?P<italic>[^*]+)\*(?!\*))"
)


def add_hyperlink(paragraph, url: str, text: str, *, code: bool = False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    if code:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rPr.append(rFonts)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, base_bold: bool = False):
    """Render a string with inline formatting into runs on the paragraph."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = base_bold
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif m.group("link_text") is not None:
            link_text = m.group("link_text")
            link_url = m.group("link_url")
            inner = INLINE_PATTERN.search(link_text)
            is_code = bool(inner and inner.group("code") is not None)
            display = inner.group("code") if is_code else link_text
            add_hyperlink(paragraph, link_url, display, code=is_code)
        elif m.group("italic") is not None:
            run = paragraph.add_run(m.group("italic"))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold


def add_horizontal_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def render_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline(p, txt, base_bold=True)
        for run in p.runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        for i, txt in enumerate(row):
            if i >= len(header):
                break
            cell = table.rows[r].cells[i]
            cell.text = ""
            add_inline(cell.paragraphs[0], txt)


def convert(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            p = doc.add_paragraph()
            add_horizontal_rule(p)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            p = doc.add_heading(level=min(level, 4))
            add_inline(p, heading_text)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|?\s*:?-{2,}", lines[i + 1].strip()):
            header = parse_table_row(stripped)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            render_table(doc, header, rows)
            continue

        # Blockquotes (single or multi-line)
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, " ".join(quote_lines))
            continue

        # Bullet lists
        if re.match(r"^[-*]\s+", stripped):
            while i < n:
                lstripped = lines[i].strip()
                if not lstripped:
                    i += 1
                    break
                m_li = re.match(r"^[-*]\s+(.*)$", lstripped)
                if not m_li:
                    break
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, m_li.group(1))
                i += 1
            continue

        # Plain paragraph (may have inline formatting). Collect contiguous lines.
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^(#{1,6})\s+", nxt):
                break
            if nxt.startswith(">") or re.match(r"^[-*]\s+", nxt):
                break
            if nxt.startswith("|"):
                break
            if re.fullmatch(r"-{3,}", nxt):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))

    doc.save(str(docx_path))


def main():
    if len(sys.argv) < 2:
        print("usage: md_to_docx.py <reports_dir>", file=sys.stderr)
        sys.exit(2)
    reports_dir = Path(sys.argv[1])
    md_files = sorted(reports_dir.glob("*.md"))
    if not md_files:
        print(f"no .md files in {reports_dir}", file=sys.stderr)
        sys.exit(1)
    for md in md_files:
        out = md.with_suffix(".docx")
        convert(md, out)
        print(f"wrote {out}")


if __name__ == "__main__":
    main()
